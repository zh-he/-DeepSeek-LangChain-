import json
import os
import time

import streamlit as st
# 注意：这里改为从 python-docx 导入 Document，以便解析本地 docx 文件
from docx import Document as DocxDocument
from langchain.docstore.document import Document
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from pdfminer.high_level import extract_text

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
VECTOR_STORE_DIR = "vector_stores"
CONVERSATION_HISTORY_DIR = "conversation_histories"
SESSIONS_FILE = "sessions.json"

VECTOR_STORE_SUBDIR = "embedding"
VECTOR_STORE_PATH = os.path.join(VECTOR_STORE_DIR, VECTOR_STORE_SUBDIR)


def get_sessions():
    """
    获取所有会话列表（存储在 sessions.json 中）。
    """
    if os.path.exists(SESSIONS_FILE):
        with open(SESSIONS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []


def save_sessions(sessions):
    """
    保存会话列表到 sessions.json。
    """
    with open(SESSIONS_FILE, 'w', encoding='utf-8') as f:
        json.dump(sessions, f, ensure_ascii=False)


def delete_session(session_name):
    """
    删除指定会话及其相关文件，并在 sessions.json 中删除该会话。
    注意：conversation_histories 下的实际存放路径是
    conversation_histories/sessions/<session_name>.json
    """

    # 1) 删除对话历史文件
    history_file = os.path.join(CONVERSATION_HISTORY_DIR, "sessions", f"{session_name}.json")
    if os.path.exists(history_file):
        os.remove(history_file)

    # 2) 从 sessions.json 中删除此会话
    sessions = get_sessions()
    if session_name in sessions:
        sessions.remove(session_name)
        save_sessions(sessions)


def load_text_file(file_path):
    """
    从纯文本文件中加载全文内容。
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            if not text.strip():
                st.warning(f"文本文件 {file_path} 中未提取到任何内容。")
            else:
                st.write(f"成功加载文本文件 {file_path}, 字符数: {len(text)}")
            return text
    except Exception as e:
        st.error(f"读取文本文件时出错: {e}")
        return ""


def load_pdf(file_path):
    """
    从 PDF 文件中提取文本。
    """
    try:
        text = extract_text(file_path)
        if not text.strip():
            st.warning(f"从 PDF 文件 {file_path} 中未提取到任何文本。")
        else:
            st.write(f"成功从 PDF 文件 {file_path} 中提取了 {len(text)} 个字符。")
        return text
    except Exception as e:
        st.error(f"从 PDF 提取文本时出错: {e}")
        return ""


def load_md(file_path):
    """
    从 Markdown 文件中读取文本。
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
            if not md_content.strip():
                st.warning(f"Markdown文件 {file_path} 中未提取到任何内容。")
            else:
                st.write(f"成功加载 Markdown 文件 {file_path}, 字符数: {len(md_content)}")
            return md_content
    except Exception as e:
        st.error(f"读取 Markdown 文件时出错: {e}")
        return ""


def load_docx(file_path):
    """
    从 Word (docx) 文档中加载文本.
    如果要兼容 .doc (老格式), 需要引入其他第三方库或先转换为 docx 格式。
    """
    try:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = "\n".join(paragraphs)
        if not full_text.strip():
            st.warning(f"Word文件 {file_path} 中未提取到任何文本。")
        else:
            st.write(f"成功从 Word (docx) 文件 {file_path} 中加载了 {len(full_text)} 个字符。")
        return full_text
    except Exception as e:
        st.error(f"加载 Word 文件时出错: {e}")
        return ""


def load_document(file_path):
    """
    根据文件扩展名加载不同类型的文档.
    支持: .pdf, .txt, .md, .docx, .doc (只警告不做处理).
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    if ext == '.pdf':
        return load_pdf(file_path)
    elif ext == '.txt':
        return load_text_file(file_path)
    elif ext == '.md':
        return load_md(file_path)
    elif ext == '.docx':
        return load_docx(file_path)
    else:
        st.error(f"不支持的文件扩展名: {ext}")
        return ""


def prepare_documents(text, chunk_size=512, chunk_overlap=64):
    """
    将文本拆分为多个文档块。
    """
    text_splitter = CharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = text_splitter.split_text(text)
    st.write(f"生成了 {len(chunks)} 个文本块。")

    if not chunks:
        st.error("文本分块后未生成任何文本块。请检查内容和分块参数。")
        return []

    documents = [Document(page_content=chunk) for chunk in chunks]
    return documents


def initialize_conversation_history(session_id):
    """初始化对话历史，根据会话ID加载对应的JSON文件"""
    os.makedirs(CONVERSATION_HISTORY_DIR, exist_ok=True)
    history_file = os.path.join(CONVERSATION_HISTORY_DIR, "sessions", f"{session_id}.json")
    if os.path.exists(history_file):
        if os.path.getsize(history_file) == 0:
            return []
        else:
            try:
                with open(history_file, 'r', encoding='utf-8') as f:
                    history = json.load(f)
                    history = [tuple(pair) for pair in history]
                    return history
            except json.JSONDecodeError:
                st.warning("对话历史文件内容无效，将初始化为空。")
                return []
    else:
        return []


def initialize_vector_store():
    """
    初始化全局的向量存储。
    如果向量存储已存在，则加载它；否则，创建一个新的向量存储。
    """
    os.makedirs(VECTOR_STORE_PATH, exist_ok=True)
    vector_store_path = VECTOR_STORE_PATH

    if os.path.exists(vector_store_path):
        try:
            embeddings = HuggingFaceEmbeddings(model_name="shibing624/text2vec-base-chinese")
            vector_store = FAISS.load_local(vector_store_path, embeddings, allow_dangerous_deserialization=True)
            st.success(f"已加载现有的向量存储：{vector_store_path}")
            return vector_store
        except Exception as e:
            st.error(f"加载向量存储时出错: {e}")
            st.info("将尝试创建一个新的向量存储。")

    # 如果向量存储不存在或加载失败，返回 None
    return None


def build_vector_store_from_documents(documents, vector_store_path):
    """
    从文档列表构建向量数据库并保存到指定路径。
    注意：vector_store_path 会被当作一个文件夹来保存 index.faiss 和 index.pkl 等文件。
    """
    # 确保目录存在
    os.makedirs(vector_store_path, exist_ok=True)

    embeddings = HuggingFaceEmbeddings(model_name="shibing624/text2vec-base-chinese")
    vector_store = FAISS.from_documents(documents, embeddings)

    if vector_store.index.ntotal == 0:
        st.error("嵌入生成失败，未生成任何向量。请检查嵌入模型和文档内容。")
        return None

    # 这里会在 vector_store_path 下生成 index.faiss 和 index.pkl
    vector_store.save_local(vector_store_path)
    st.success(f"向量数据库已保存到 {vector_store_path}")
    return vector_store


def add_documents_to_vector_store(vector_store, documents):
    """
    将新文档添加到现有向量数据库（不需要新建）。
    """
    try:
        vector_store.add_documents(documents)
        vector_store.save_local(VECTOR_STORE_PATH)  # 确保保存更新后的向量存储
        st.write(f"已添加 {len(documents)} 个文档到向量数据库，并保存更新。")
    except Exception as e:
        st.error(f"添加文档到向量存储时出错: {e}")


def save_conversation_history(session_id, history):
    """
    保存对话历史到 conversation_histories/sessions/<session_id>.json
    """
    os.makedirs(os.path.join(CONVERSATION_HISTORY_DIR, "sessions"), exist_ok=True)
    history_file = os.path.join(CONVERSATION_HISTORY_DIR, "sessions", f"{session_id}.json")
    with open(history_file, 'w', encoding='utf-8') as f:
        json.dump([list(pair) for pair in history], f, ensure_ascii=False, indent=4)


def conversational_answer(chain, question, stop_flag, history):
    """
    通用的多轮对话回答流程，返回最终 answer 和 source_documents。
    支持 ConversationalRetrievalChain 和普通的 LLMChain
    """
    try:
        # 检查是否是 ConversationalRetrievalChain（有文档检索功能的 chain）
        if hasattr(chain, 'combine_docs_chain'):
            # 使用带文档的 chain
            for _ in range(3):  # 保留模拟生成效果
                if stop_flag["stop"]:
                    return {"answer": "回答已终止。", "source_documents": []}
                time.sleep(1)

            # 调用 chain 时使用正确的参数格式
            response = chain({"question": question, "chat_history": history})
            return {
                "answer": response.get("answer", ""),
                "source_documents": response.get("source_documents", [])
            }
        else:
            # 使用普通的 LLMChain（如 fallback chain）
            for _ in range(3):  # 保留模拟生成效果
                if stop_flag["stop"]:
                    return {"answer": "回答已终止。", "source_documents": []}
                time.sleep(1)

            # 对于普通的 LLMChain，使用 run 方法
            response = chain.run(question=question, chat_history=history)

            # LLMChain 返回的是字符串
            return {
                "answer": response if isinstance(response, str) else str(response),
                "source_documents": []
            }

    except Exception as e:
        st.error(f"在生成回答时出错: {e}")
        return {"answer": "抱歉，我无法生成回答。", "source_documents": []}
